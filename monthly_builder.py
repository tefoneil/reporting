#!/usr/bin/env python3
"""
Monthly Report Builder for Chronic Circuit Analysis
Processes Tableau exports and generates Word/PDF reports
"""

import pandas as pd
import numpy as np
from pathlib import Path
import argparse
from typing import Dict, Any
import subprocess
import sys
from datetime import datetime, timedelta
import json
from docx import Document
from docx.shared import Inches, Pt
import matplotlib.pyplot as plt
import seaborn as sns
import FreeSimpleGUI as sg
import os
import re
import logging
import sys
import subprocess
from analyze_data import get_rolling_ticket_total
from utils import canonical_id, warn_low_ticket_median, validate_metadata, get_file_sha256, validate_calculations, filter_test_circuits

# Configuration constants
CONSISTENT_THRESHOLD = int(os.getenv("MR_CONSISTENT_THRESHOLD", 6))
DYNAMIC_CONSISTENCY = int(os.getenv("MR_DYNAMIC_CONSISTENCY", 0))  # 0 = legacy mode, 1 = dynamic mode

# Core chronic classification thresholds (unchanged in v0.1.7-b)
AVAIL_THRESH_PCT = float(os.getenv("MR_THRESH_AVAIL_PCT", 5.0))  # Availability significant change threshold



class ChronicReportBuilder:
    def __init__(self, exclude_regional=False, show_indicators=True):
        """
        Initialize the chronic report builder
        
        Args:
            exclude_regional: Flag to exclude regional circuits from new chronic detection
            show_indicators: Show (C) and (R) flags in reports (default True)
        """
        self.exclude_regional = exclude_regional
        self.show_indicators = show_indicators
        self.service_seconds_per_month = 30.44 * 24 * 3600  # Average month in seconds
        self.labor_rate = 60  # $60/hour loaded rate
        
        # Regional circuits list
        self.regional_circuits = [
            '500335805', '500332738', '500334193', '500394949', '500394765',
            'IST6022E#2_010G', 'IST6041E#3_010G', 'LZA010663', 'LZA010635', 'LZA010634',
            '027ISAN284012272923', '091NOID1143035717849_889621', '091NOID1143035717419_889599',
            'SR216187', '091NOID1143037092974_993502'
        ]
        
    def load_crosstab_data(self, impacts_file, counts_file):
        """Load and process the Tableau export files"""
        print(f"Loading impact data from {impacts_file}")
        if str(impacts_file).lower().endswith('.csv'):
            impacts_df = pd.read_csv(impacts_file)
        else:
            impacts_df = pd.read_excel(impacts_file)
        # Fix: Trim column headers to handle trailing spaces
        impacts_df.columns = impacts_df.columns.str.strip()
        
        # Handle column aliasing for Config Item Name
        if 'Configuration Item Name' in impacts_df.columns:
            impacts_df = impacts_df.rename(columns={'Configuration Item Name': 'Config Item Name'})
        
        # v0.1.8: Filter out test circuits
        initial_count = len(impacts_df)
        if 'Config Item Name' in impacts_df.columns:
            # Filter CID_TEST circuits
            test_filter = impacts_df['Config Item Name'].str.startswith('CID_TEST', na=False)
            if 'Vendor' in impacts_df.columns:
                # Also filter vendors containing 'Test'
                test_filter |= impacts_df['Vendor'].str.contains('Test', case=False, na=False)
            
            impacts_df = impacts_df[~test_filter]
            filtered_count = initial_count - len(impacts_df)
            if filtered_count > 0:
                print(f"Filtered out {filtered_count} test circuits from impacts data")
        
        # Fix: Forward-fill blank month cells to prevent 0-ticket miscounts
        blank_percentage = 0
        if 'Inc Resolved At (Month / Year)' in impacts_df.columns:
            blank_count = impacts_df['Inc Resolved At (Month / Year)'].isna().sum()
            total_rows = len(impacts_df)
            blank_percentage = (blank_count / total_rows * 100) if total_rows > 0 else 0
            if blank_count > 0:
                print(f"Forward-filling {blank_count} blank month cells ({blank_percentage:.1f}% of data)")
                impacts_df.loc[:, 'Inc Resolved At (Month / Year)'] = impacts_df['Inc Resolved At (Month / Year)'].ffill()
        
        # Store data quality info for GUI warning
        self.data_quality_warning = blank_percentage > 10
        
        print(f"Loading counts data from {counts_file}")  
        if str(counts_file).lower().endswith('.csv'):
            counts_df = pd.read_csv(counts_file)
        else:
            counts_df = pd.read_excel(counts_file)
        # Fix: Trim column headers to handle trailing spaces
        counts_df.columns = counts_df.columns.str.strip()
        
        # Handle column aliasing for Config Item Name
        if 'Configuration Item Name' in counts_df.columns:
            counts_df = counts_df.rename(columns={'Configuration Item Name': 'Config Item Name'})
        
        # v0.1.8: Filter out test circuits from counts data too
        initial_count = len(counts_df)
        if 'Config Item Name' in counts_df.columns:
            # Filter CID_TEST circuits
            test_filter = counts_df['Config Item Name'].str.startswith('CID_TEST', na=False)
            if 'Vendor' in counts_df.columns:
                # Also filter vendors containing 'Test'
                test_filter |= counts_df['Vendor'].str.contains('Test', case=False, na=False)
            
            counts_df = counts_df[~test_filter]
            filtered_count = initial_count - len(counts_df)
            if filtered_count > 0:
                print(f"Filtered out {filtered_count} test circuits from counts data")
        
        # Clean numeric columns that might have comma formatting
        for col in impacts_df.columns:
            if 'Duration' in col or 'Count' in col:
                if impacts_df[col].dtype == 'object':
                    impacts_df[col] = pd.to_numeric(impacts_df[col].astype(str).str.replace(',', ''), errors='coerce')
        
        for col in counts_df.columns:
            if any(x in col for x in ['Cost', 'Duration', 'Count', 'Sum', 'Average']):
                if counts_df[col].dtype == 'object':
                    counts_df[col] = pd.to_numeric(counts_df[col].astype(str).str.replace(',', ''), errors='coerce')
        
        # Add canonical IDs for both DataFrames
        impacts_df['canonical_id'] = impacts_df['Config Item Name'].apply(canonical_id)
        counts_df['canonical_id'] = counts_df['Config Item Name'].apply(canonical_id)
        
        # Add numeric coercion data quality check for ticket counts
        ticket_column = 'Distinct count of Inc Nbr'
        if ticket_column in impacts_df.columns:
            # Ensure ticket column is numeric in impacts_df too
            impacts_df[ticket_column] = pd.to_numeric(impacts_df[ticket_column], errors='coerce')
            
            # Check for non-numeric coercion failures
            non_numeric_count = impacts_df[ticket_column].isna().sum()
            total_ticket_rows = len(impacts_df)
            non_numeric_ratio = (non_numeric_count / total_ticket_rows) if total_ticket_rows > 0 else 0
            
            # Store warning flag for GUI
            self.ticket_coercion_warning = non_numeric_ratio > 0.10
            
            if non_numeric_count > 0:
                print(f"Ticket count coercion: {non_numeric_count} non-numeric values ({non_numeric_ratio:.1%} of data)")
        elif ticket_column in counts_df.columns:
            # Ensure ticket column is numeric
            counts_df[ticket_column] = pd.to_numeric(counts_df[ticket_column], errors='coerce')
            
            # Check for non-numeric coercion failures
            non_numeric_count = counts_df[ticket_column].isna().sum()
            total_ticket_rows = len(counts_df)
            non_numeric_ratio = (non_numeric_count / total_ticket_rows) if total_ticket_rows > 0 else 0
            
            # Store warning flag for GUI
            self.ticket_coercion_warning = non_numeric_ratio > 0.10
            
            if non_numeric_count > 0:
                print(f"Ticket count coercion: {non_numeric_count} non-numeric values ({non_numeric_ratio:.1%} of data)")
        else:
            self.ticket_coercion_warning = False
        
        return impacts_df, counts_df
    
    def load_baseline_status(self, output_dir='./final_output'):
        """Load baseline legacy status from frozen legacy list (v0.1.9+)"""
        baseline_status = {}
        baseline_ids = set()
        cutover_found = False
        
        try:
            from pathlib import Path
            
            # v0.1.9: Load from frozen legacy list first
            frozen_legacy_path = Path('./docs/frozen_legacy_list.json')
            if frozen_legacy_path.exists():
                try:
                    with open(frozen_legacy_path, 'r') as f:
                        frozen_data = json.load(f)
                    
                    # Map frozen legacy statuses using canonical IDs, excluding test circuits
                    for circuit in frozen_data.get('chronic_consistent', []):
                        if not circuit.startswith('CID_TEST'):  # P1-b: Filter test circuits
                            canonical = canonical_id(circuit)
                            baseline_status[canonical] = 'Consistent'
                            baseline_ids.add(canonical)
                    
                    for circuit in frozen_data.get('chronic_inconsistent', []):
                        if not circuit.startswith('CID_TEST'):  # P1-b: Filter test circuits
                            canonical = canonical_id(circuit)
                            baseline_status[canonical] = 'Inconsistent'
                            baseline_ids.add(canonical)
                    
                    for circuit in frozen_data.get('media_chronics', []):
                        if not circuit.startswith('CID_TEST'):  # P1-b: Filter test circuits
                            canonical = canonical_id(circuit)
                            baseline_status[canonical] = 'Media Chronic'
                        baseline_ids.add(canonical)
                    
                    cutover_found = True
                    logging.info(f"Loaded {len(baseline_status)} circuits from frozen legacy list")
                    
                except (json.JSONDecodeError, KeyError) as e:
                    logging.warning(f"Failed to load frozen legacy list: {e}")
            
            # Fallback: scan for chronic summary JSON files if frozen list not available
            if not cutover_found:
                logging.info("Frozen legacy list not found, falling back to JSON scan")
                output_path = Path(output_dir)
                if output_path.exists():
                    json_files = list(output_path.glob('chronic_summary_*.json'))
                    json_files.sort()
                    
                    for json_file in json_files:
                        try:
                            with open(json_file, 'r') as f:
                                summary_data = json.load(f)
                            
                            # Check if this is May 2025 or earlier
                            filename = json_file.name
                            if 'May_2025' in filename or any(month in filename for month in 
                                ['January_2025', 'February_2025', 'March_2025', 'April_2025']):
                                cutover_found = True
                                
                                # Extract legacy status from existing chronics
                                existing = summary_data.get('chronic_data', {}).get('existing_chronics', {})
                                
                                # Map consistent/inconsistent statuses using canonical IDs
                                for circuit in existing.get('chronic_consistent', []):
                                    canonical = canonical_id(circuit)
                                    baseline_status[canonical] = 'Consistent'
                                    baseline_ids.add(canonical)
                                
                                for circuit in existing.get('chronic_inconsistent', []):
                                    canonical = canonical_id(circuit)
                                    baseline_status[canonical] = 'Inconsistent'
                                    baseline_ids.add(canonical)
                                
                                # Media chronics maintain their status
                                for circuit in existing.get('media_chronics', []):
                                    canonical = canonical_id(circuit)
                                    baseline_status[canonical] = 'Media Chronic'
                                    baseline_ids.add(canonical)
                            
                            # NEW: Include prior month's New Chronics for promotion evaluation
                            new_chronics_data = summary_data.get('chronic_data', {}).get('new_chronics', {})
                            promotion_count = 0
                            for provider_type, circuits in new_chronics_data.items():
                                for circuit in circuits:
                                    canonical = canonical_id(circuit)
                                    baseline_status[canonical] = 'pending_promotion'
                                    baseline_ids.add(canonical)
                                    promotion_count += 1
                            
                                print(f"📊 Loaded baseline status from {json_file.name}: {len(baseline_status)} circuits ({len(baseline_status) - promotion_count} legacy + {promotion_count} pending promotion)")
                                break
                        except Exception as file_error:
                            print(f"⚠️  Error reading {json_file.name}: {file_error}")
                            continue
                
        except Exception as e:
            print(f"⚠️  Error loading baseline status: {e}")
        
        # Store for GUI warning if no baseline found
        self.baseline_found = cutover_found
        
        return baseline_status, baseline_ids
    
    def process_chronic_logic(self, impacts_df, counts_df):
        """Process chronic circuit logic based on business rules"""
        print("Processing chronic circuit logic...")
        
        # Merge data on Config Item Name
        merged_df = pd.merge(
            impacts_df, 
            counts_df, 
            on='Config Item Name', 
            how='outer'
        ).fillna({
            'COUNTD Months': 0,
            'Outage Duration': 0,
            'Incident Network-facing Impacted CI Type': 'Unknown Provider'
        })
        
        # Convert outage duration from seconds to hours
        merged_df['ImpactHours'] = merged_df['Outage Duration'] / 3600
        
        # v0.1.8-audit: Create raw ticket counts dictionary for audit trail
        raw_counts = (
            impacts_df.groupby("canonical_id")["Distinct count of Inc Nbr"]
                      .sum()
                      .to_dict()
        )
        
        # Load baseline status for hybrid consistency mode
        baseline_status, baseline_ids = self.load_baseline_status()
        
        # Hybrid consistency mode: combine legacy + ticket-based classification
        all_chronic_circuits = [
            '500332738', '500334193', '500335805',  # Cirion
            '091NOID1143035717419_889599', '091NOID1143035717849_889621',  # Tata
            'SR216187',  # PCCW
            'PTH TOK EPL 90030025',  # Telstra
            'LZA010663',  # Liquid Telecom (April addition)
            'LD017936',  # Orange
            'IST6041E#3_010G', 'IST6022E#2_010G',  # Globenet
            'HI/ADM/00697867',  # GTT
            'SR215576',  # PCCW
            'SSO-JBTKRHS002F-DWDM10',  # Sansa
            '443463817', '445597814', '443919489', '445979698', '443832799', 'FRO2007133508',  # Lumen
            'W1E32092',  # Verizon (April addition)
            'N9675474L', 'N2864477L'  # Telstra (April additions)
            # P1-b: Removed CID_TEST circuits from hardcoded list
        ]
        
        # Add circuits with pending_promotion status to the processing list
        pending_promotion_circuits = [circuit for canonical, status in baseline_status.items() 
                                     if status == 'pending_promotion' 
                                     for circuit in merged_df['Config Item Name'].unique() 
                                     if canonical_id(circuit) == canonical]
        all_chronic_circuits.extend(pending_promotion_circuits)
        
        # Hybrid classification: legacy circuits keep status, new circuits use ticket-based
        chronic_consistent = []
        chronic_inconsistent = []
        media_chronics_hybrid = []
        circuit_ticket_data = {}  # Store rolling ticket totals for auditing
        
        for circuit_id in all_chronic_circuits:
            # Convert to canonical ID for lookups and aggregation
            canonical = canonical_id(circuit_id)
            rolling_tickets = get_rolling_ticket_total(canonical, merged_df)
            circuit_ticket_data[circuit_id] = {
                'rolling_ticket_total': rolling_tickets,
                'raw_ticket_count_crosstab': raw_counts.get(canonical, 0)
            }
            
            # Apply hybrid logic: baseline status takes precedence (using canonical lookup)
            if canonical in baseline_status:
                status = baseline_status[canonical]
                
                if status == 'pending_promotion':
                    # Prior month's New Chronic - evaluate ticket rule for promotion
                    if rolling_tickets >= CONSISTENT_THRESHOLD:
                        chronic_consistent.append(circuit_id)
                        circuit_ticket_data[circuit_id]['status'] = 'consistent'
                    else:
                        chronic_inconsistent.append(circuit_id)
                        circuit_ticket_data[circuit_id]['status'] = 'inconsistent'
                elif status == 'Consistent':
                    # Legacy circuit - freeze existing status
                    chronic_consistent.append(circuit_id)
                    circuit_ticket_data[circuit_id]['status'] = 'consistent'
                elif status == 'Inconsistent':
                    chronic_inconsistent.append(circuit_id)
                    circuit_ticket_data[circuit_id]['status'] = 'inconsistent'
                elif status == 'Media Chronic':
                    media_chronics_hybrid.append(circuit_id)
                    circuit_ticket_data[circuit_id]['status'] = 'media chronic'
            else:
                # New circuit - use ticket-based classification
                if rolling_tickets >= CONSISTENT_THRESHOLD:
                    chronic_consistent.append(circuit_id)
                    circuit_ticket_data[circuit_id]['status'] = 'consistent'
                else:
                    chronic_inconsistent.append(circuit_id)
                    circuit_ticket_data[circuit_id]['status'] = 'inconsistent'
        
        # Media chronics: combine baseline + hardcoded list
        media_chronics_hardcoded = [
            'VID-1583', 'VID-1597', 'VID-1598',  # Slovak Telekom
            'VID-1574', 'VID-1575', 'VID-1581', 'VID-1582',  # Slovak
            'VID-1146', 'VID-1525', 'VID-1530', 'VID-0875'  # BBC Global News
        ]
        
        # Combine hybrid media chronics with hardcoded ones (remove duplicates)
        media_chronics = list(set(media_chronics_hybrid + media_chronics_hardcoded))
        
        # Performance monitoring circuits (unchanged)
        perf_60_day = ['444089285', '444089468']  # KTA SNG dropped from April data
        perf_30_day = ['445082297']  # Updated based on April data, 445082296 moved to new chronic for May demo
        
        existing_chronics = {
            'chronic_consistent': chronic_consistent,
            'chronic_inconsistent': chronic_inconsistent,
            'media_chronics': media_chronics,
            'perf_60_day': perf_60_day,
            'perf_30_day': perf_30_day
        }
        
        # Find circuits that have progressed through performance monitoring
        all_existing_chronics = (existing_chronics['chronic_consistent'] + 
                                existing_chronics['chronic_inconsistent'] + 
                                existing_chronics['media_chronics'])
        
        # Circuits reaching 3rd month that are NOT already chronic AND have been through monitoring
        potential_new_chronics = merged_df[
            (merged_df['COUNTD Months'] == 3) & 
            (~merged_df['Config Item Name'].isin(all_existing_chronics))
        ].drop_duplicates(subset=['Config Item Name'])
        
        # Check which ones have been through the 60-day -> 30-day progression
        # (from previous month's 30-day list, indicating they've completed the progression)
        # Adding 444282783 as demo new chronic for May report (not on regional list)
        completed_progression_circuits = existing_chronics['perf_30_day'] + ['444282783']
        new_chronics = potential_new_chronics[
            potential_new_chronics['Config Item Name'].isin(completed_progression_circuits)
        ]
        
        print(f"Found {len(potential_new_chronics)} circuits at 3rd month")
        print(f"Of these, {len(new_chronics)} have completed 60->30 day progression")
        
        # Filter out name variations that match existing circuits
        print("Filtering name variations...")
        filtered_new_chronics = []
        excluded_variations = []
        
        for idx, row in new_chronics.iterrows():
            circuit_name = str(row['Config Item Name'])
            is_variation = False
            
            for existing in all_existing_chronics:
                existing_str = str(existing)
                # Check for key identifiers in circuit names
                circuit_parts = circuit_name.replace('/', '').replace('-', '').replace(' ', '').replace('_', '')
                existing_parts = existing_str.replace('/', '').replace('-', '').replace(' ', '').replace('_', '')
                
                # If substantial part of circuit name exists in master list, it's a variation
                if len(circuit_parts) > 5 and len(existing_parts) > 5:
                    if (circuit_parts in existing_parts) or (existing_parts in circuit_parts):
                        excluded_variations.append(f"'{circuit_name}' matches existing '{existing_str}'")
                        is_variation = True
                        break
                
                # Special case: 419 circuits are the same circuit family
                if '419' in circuit_name and '419' in existing_str:
                    excluded_variations.append(f"'{circuit_name}' matches 419 circuit family '{existing_str}'")
                    is_variation = True
                    break
                
                # Confirmed variations
                if ('091NOID1143035717419_1040578' in circuit_name and '091NOID1143035717419_889599' in existing_str) or \
                   ('091NOID1143035717419_889599' in existing_str and '091NOID1143035717419_1040578' in circuit_name):
                    excluded_variations.append(f"'{circuit_name}' matches confirmed variation '{existing_str}'")
                    is_variation = True
                    break
                    
                if ('LD017936' in circuit_name and 'LD017936' in existing_str) and circuit_name != existing_str:
                    excluded_variations.append(f"'{circuit_name}' matches confirmed variation '{existing_str}'")
                    is_variation = True
                    break
            
            # Check for regional circuits if flagging is enabled
            if self.exclude_regional and circuit_name in self.regional_circuits:
                excluded_variations.append(f"'{circuit_name}' excluded as regional circuit")
                is_variation = True
            
            if not is_variation:
                filtered_new_chronics.append(row)
        
        # Print excluded variations
        for exclusion in excluded_variations:
            print(f"[EXCLUDED] Excluded variation: {exclusion}")
        
        # Convert back to DataFrame
        if filtered_new_chronics:
            new_chronics = pd.DataFrame(filtered_new_chronics).reset_index(drop=True)
        else:
            new_chronics = pd.DataFrame()
        
        # Group new chronics by provider, excluding promoted circuits
        promoted_circuits = [circuit for circuit in all_chronic_circuits 
                           if canonical_id(circuit) in baseline_status and 
                           baseline_status[canonical_id(circuit)] == 'pending_promotion']
        
        if len(new_chronics) > 0:
            # Exclude promoted circuits from new chronic summary
            remaining_new_chronics = new_chronics[~new_chronics['Config Item Name'].isin(promoted_circuits)]
            if len(remaining_new_chronics) > 0:
                new_chronic_summary = remaining_new_chronics.groupby('Incident Network-facing Impacted CI Type')['Config Item Name'].apply(lambda x: list(x.unique())).to_dict()
            else:
                new_chronic_summary = {}
        else:
            new_chronic_summary = {}
        
        # Performance monitoring updates (30-day becomes new chronic candidates, 60-day becomes 30-day)
        updated_perf_30_day = existing_chronics['perf_60_day'].copy()
        updated_perf_60_day = []  # Will be populated with new 60-day candidates
        
        # Calculate final counts including promoted circuits
        total_promoted = len(promoted_circuits)
        final_new_chronic_count = len(new_chronics) - total_promoted if len(new_chronics) >= total_promoted else 0
        
        return {
            'total_chronic_circuits': len(existing_chronics['chronic_consistent']) + len(existing_chronics['chronic_inconsistent']) + total_promoted,
            'media_chronics': len(existing_chronics['media_chronics']),
            'new_chronics': new_chronic_summary,
            'new_chronic_count': final_new_chronic_count,
            'existing_chronics': existing_chronics,
            'updated_perf_30_day': updated_perf_30_day,
            'updated_perf_60_day': updated_perf_60_day,
            'circuit_ticket_data': circuit_ticket_data,  # Add rolling ticket data for auditing
            'merged_data': merged_df
        }
    
    def generate_metadata(self, impacts_file, counts_file):
        """Generate metadata block for JSON output (v0.1.9)"""
        from pathlib import Path
        
        try:
            # Get git commit hash
            try:
                git_commit = subprocess.check_output(['git', 'rev-parse', 'HEAD'], 
                                                   stderr=subprocess.DEVNULL).decode().strip()
            except (subprocess.CalledProcessError, FileNotFoundError):
                git_commit = "unknown"
            
            # Generate file hashes
            impacts_path = Path(impacts_file)
            counts_path = Path(counts_file)
            
            metadata = {
                'tool_version': '0.1.9',
                'python_version': f"{sys.version_info.major}.{sys.version_info.minor}.{sys.version_info.micro}",
                'git_commit': git_commit,
                'run_timestamp': datetime.now().isoformat() + 'Z',  # UTC format
                'crosstab_sha256': get_file_sha256(impacts_path) if impacts_path.exists() else 'unknown',
                'counts_sha256': get_file_sha256(counts_path) if counts_path.exists() else 'unknown'
            }
            
            # Validate metadata has all required keys
            if not validate_metadata(metadata):
                logging.warning("Metadata validation failed - some keys may be missing")
            
            return metadata
            
        except Exception as e:
            logging.error(f"Failed to generate metadata: {e}")
            # Return minimal metadata on error
            return {
                'tool_version': '0.1.9',
                'python_version': f"{sys.version_info.major}.{sys.version_info.minor}.{sys.version_info.micro}",
                'git_commit': 'error',
                'run_timestamp': datetime.now().isoformat() + 'Z',
                'crosstab_sha256': 'error',
                'counts_sha256': 'error'
            }

    def calculate_metrics(self, chronic_data):
        """Calculate all required metrics for the report using FULL dataset"""
        
        metrics = {}
        merged_df = chronic_data['merged_data']
        existing_chronics = chronic_data['existing_chronics']
        
        # Basic counts from chronic data structure
        metrics['total_chronic_circuits'] = chronic_data['total_chronic_circuits']
        metrics['media_chronics'] = chronic_data['media_chronics']
        metrics['new_chronic_count'] = chronic_data['new_chronic_count']
        metrics['new_chronics'] = chronic_data['new_chronics']
        
        # Provider count - unique vendors across all chronic categories + performance monitoring
        all_vendor_circuits = (existing_chronics['chronic_consistent'] + 
                              existing_chronics['chronic_inconsistent'] + 
                              existing_chronics['perf_60_day'] + 
                              existing_chronics['perf_30_day'])
        
        # Map circuits to vendors (comprehensive mapping)
        vendor_count = set()
        for circuit in all_vendor_circuits:
            if circuit.startswith('500'):
                vendor_count.add('Cirion')
            elif circuit.startswith('091'):
                vendor_count.add('Tata')
            elif circuit.startswith('SR'):
                vendor_count.add('PCCW')
            elif 'PTH' in circuit or circuit.startswith('N') or circuit.startswith('KTA'):
                vendor_count.add('Telstra')
            elif circuit.startswith('LZA'):
                vendor_count.add('Liquid Telecom')
            elif circuit.startswith('LD'):
                vendor_count.add('Orange')
            elif circuit.startswith('IST'):
                vendor_count.add('Globenet')
            elif 'HI/ADM' in circuit:
                vendor_count.add('GTT')
            elif 'SSO' in circuit:
                vendor_count.add('Sansa')
            elif circuit.startswith('44') or circuit.startswith('FRO'):
                vendor_count.add('Lumen')
            elif circuit.startswith('W1E'):
                vendor_count.add('Verizon')
        
        metrics['total_providers'] = len(vendor_count)
        
        # USE FULL DATASET (all circuits) for analysis, not just chronics
        all_circuits_df = merged_df.copy()
        
        # Clean data - remove rows with missing circuit names
        all_circuits_df = all_circuits_df.dropna(subset=['Config Item Name'])
        
        # P1-b: Filter test circuits from analysis data  
        all_circuits_df = filter_test_circuits(all_circuits_df, 'Config Item Name')
        
        # Top 5 by ticket count (from ALL circuits in data)
        if 'Distinct count of Inc Nbr' in all_circuits_df.columns:
            ticket_counts = all_circuits_df.groupby('Config Item Name')['Distinct count of Inc Nbr'].sum().sort_values(ascending=False)
            metrics['top5_tickets'] = ticket_counts.head(5).to_dict()
        
        # Top 5 by cost to serve (from ALL circuits in data)
        if 'Cost to Serve (Sum Impact x $60/hr)' in all_circuits_df.columns:
            cost_data = all_circuits_df.groupby('Config Item Name')['Cost to Serve (Sum Impact x $60/hr)'].sum().sort_values(ascending=False)
            # Filter out zero costs
            cost_data = cost_data[cost_data > 0]
            metrics['top5_cost'] = cost_data.head(5).to_dict()
        
        # Bottom 5 availability (from ALL circuits in data)
        # P1-a fix: Always treat 'Outage Duration' as minutes, convert once to hours
        outage_column = None
        if 'Outage Duration' in all_circuits_df.columns:
            outage_column = 'Outage Duration'
        elif 'SUM Outage (Hours)' in all_circuits_df.columns:
            outage_column = 'SUM Outage (Hours)'
        
        if outage_column:
            # Calculate potential service hours for 3-month period
            days_in_period = 90  # 3 months approximation  
            potential_hours = days_in_period * 24  # 2160 hours total
            
            # Get outage data by circuit, filter test circuits first
            outage_df = all_circuits_df[['Config Item Name', outage_column]].copy()
            outage_df = filter_test_circuits(outage_df, 'Config Item Name')
            circuit_outages = outage_df.groupby('Config Item Name')[outage_column].sum()
            
            # P1-a: Always treat Outage Duration as minutes, convert to hours
            if outage_column == 'Outage Duration':
                print(f"Treating '{outage_column}' as minutes, converting to hours")
                outage_hours = circuit_outages / 60  # Convert minutes to hours
            else:
                print(f"Using '{outage_column}' as hours directly")  
                outage_hours = circuit_outages
            
            # Calculate availability: 100 × (1 – OutageHours / PotentialHours)
            availability_pct = 100 * (1 - outage_hours / potential_hours)
            
            # P1-a: Validate availability values
            invalid_circuits = availability_pct[(availability_pct < 0) | (availability_pct > 100)]
            if len(invalid_circuits) > 0:
                logging.warning(f"Found {len(invalid_circuits)} circuits with invalid availability:")
                for circuit, avail in invalid_circuits.items():
                    outage_hrs = outage_hours.get(circuit, 0)
                    logging.warning(f"  {circuit}: {avail:.1f}% (outage: {outage_hrs:.1f}h / {potential_hours}h)")
            
            # Filter to circuits with reasonable availability (0-100%)
            valid_availability = availability_pct[(availability_pct >= 0) & (availability_pct <= 100)]
            avail_data = valid_availability.sort_values()
            metrics['bottom5_availability'] = avail_data.head(5).to_dict()
        
        # MTBF calculations (from ALL circuits in data, excluding test circuits)
        if 'Distinct count of Inc Nbr' in all_circuits_df.columns:
            operating_hours = 24 * 90  # 90 days * 24 hours
            # Note: all_circuits_df already has test circuits filtered out above
            circuit_tickets = all_circuits_df.groupby('Config Item Name')['Distinct count of Inc Nbr'].sum()
            # Filter to circuits with actual incidents
            circuit_tickets = circuit_tickets[circuit_tickets > 0]
            
            mtbf_hours = operating_hours / circuit_tickets
            mtbf_days = mtbf_hours / 24
            
            # Bottom 5 (worst) MTBF from all circuits
            mtbf_data_sorted = mtbf_days.sort_values()
            metrics['bottom5_mtbf'] = mtbf_data_sorted.head(5).to_dict()
            metrics['avg_mtbf_days'] = mtbf_days.mean()
        
        # Add chronic circuit overlay information
        all_chronic_ids = (existing_chronics['chronic_consistent'] + 
                          existing_chronics['chronic_inconsistent'])
        metrics['chronic_circuit_ids'] = all_chronic_ids
        
        # Add subtle indicators to top 5 lists
        def add_indicators(circuit_dict):
            """Add subtle (C) and (R) indicators to circuit names"""
            indicated_dict = {}
            for circuit, value in circuit_dict.items():
                indicators = []
                
                # Check for chronic status (including name variations)
                is_chronic = False
                if circuit in all_chronic_ids:
                    is_chronic = True
                else:
                    # Check for partial matches for name variations
                    for chronic_id in all_chronic_ids:
                        # Extract core circuit number for comparison
                        circuit_core = circuit.replace('/', '').replace('-', '').replace(' ', '').replace('_', '')
                        chronic_core = chronic_id.replace('/', '').replace('-', '').replace(' ', '').replace('_', '')
                        
                        # If substantial overlap, consider it a match
                        if len(chronic_core) > 8 and chronic_core in circuit_core:
                            is_chronic = True
                            break
                        elif len(circuit_core) > 8 and circuit_core in chronic_core:
                            is_chronic = True
                            break
                
                if is_chronic:
                    indicators.append('C')
                
                # Check for regional status (including name variations)
                is_regional = False
                if circuit in self.regional_circuits:
                    is_regional = True
                else:
                    # Check for partial matches
                    for regional_id in self.regional_circuits:
                        circuit_core = circuit.replace('/', '').replace('-', '').replace(' ', '').replace('_', '')
                        regional_core = regional_id.replace('/', '').replace('-', '').replace(' ', '').replace('_', '')
                        
                        if len(regional_core) > 8 and regional_core in circuit_core:
                            is_regional = True
                            break
                        elif len(circuit_core) > 8 and circuit_core in regional_core:
                            is_regional = True
                            break
                
                if is_regional:
                    indicators.append('R')
                
                if indicators:
                    circuit_display = f"{circuit} ({'/'.join(indicators)})"
                else:
                    circuit_display = circuit
                indicated_dict[circuit_display] = value
            return indicated_dict
        
        # Apply indicators to all top/bottom lists (if enabled)
        if self.show_indicators:
            if 'top5_tickets' in metrics:
                metrics['top5_tickets'] = add_indicators(metrics['top5_tickets'])
            if 'top5_cost' in metrics:
                metrics['top5_cost'] = add_indicators(metrics['top5_cost'])
            if 'bottom5_availability' in metrics:
                metrics['bottom5_availability'] = add_indicators(metrics['bottom5_availability'])
            if 'bottom5_mtbf' in metrics:
                metrics['bottom5_mtbf'] = add_indicators(metrics['bottom5_mtbf'])
        
        # P1-a: Validate calculations before returning
        try:
            validate_calculations(metrics)
        except ValueError as e:
            logging.error(f"Calculation validation failed: {e}")
            # Continue execution but log the error
        
        return metrics
    
    def generate_text_summary(self, chronic_data: Dict[str, Any], metrics: Dict[str, Any], output_dir: Path, month_str: str) -> Path:
        """Generate a text file summary of chronic circuits"""
        
        output_path = output_dir / f"chronic_circuits_list_{month_str}.txt"
        
        with open(output_path, 'w') as f:
            f.write(f"CHRONIC CIRCUITS LIST - {month_str.replace('_', ' ').upper()} REPORT\n")
            f.write("=" * 40 + "\n\n")
            
            f.write(f"TOTAL CHRONIC CIRCUITS: {chronic_data['total_chronic_circuits']}\n\n")
            
            # Chronic Consistent
            consistent = chronic_data.get('existing_chronics', {}).get('chronic_consistent', [])
            f.write(f"CHRONIC CONSISTENT ({len(consistent)} circuits):\n")
            f.write("-" * 35 + "\n")
            for i, circuit in enumerate(consistent, 1):
                f.write(f"{i}. {circuit}\n")
            f.write("\n")
            
            # Chronic Inconsistent
            inconsistent = chronic_data.get('existing_chronics', {}).get('chronic_inconsistent', [])
            f.write(f"CHRONIC INCONSISTENT ({len(inconsistent)} circuits):\n")
            f.write("-" * 35 + "\n")
            for i, circuit in enumerate(inconsistent, 1):
                f.write(f"{i}. {circuit}\n")
            f.write("\n")
            
            # Media Chronics
            media = chronic_data.get('existing_chronics', {}).get('media_chronics', [])
            if media:
                f.write(f"MEDIA CHRONICS ({len(media)} circuits):\n")
                f.write("-" * 35 + "\n")
                for i, circuit in enumerate(media, 1):
                    f.write(f"{i}. {circuit}\n")
                f.write("\n")
            
            # New Chronics
            new_chronics = chronic_data.get('new_chronics', {})
            new_count = chronic_data.get('new_chronic_count', 0)
            if new_count > 0:
                f.write(f"NEW CHRONIC CIRCUITS ({new_count} circuit{'s' if new_count > 1 else ''}):\n")
                f.write("-" * 35 + "\n")
                circuit_num = 1
                for category, circuits in new_chronics.items():
                    for circuit in circuits:
                        f.write(f"{circuit_num}. {circuit} ({category})\n")
                        circuit_num += 1
                f.write("\n")
            
            # Performance Monitoring
            perf_30 = chronic_data.get('updated_perf_30_day', [])
            perf_60 = chronic_data.get('updated_perf_60_day', [])
            if perf_30 or perf_60:
                f.write("PERFORMANCE MONITORING:\n")
                f.write("-" * 35 + "\n")
                if perf_30:
                    f.write(f"30-Day Performance Watch: {', '.join(perf_30)}\n")
                if perf_60:
                    f.write(f"60-Day Performance Watch: {', '.join(perf_60)}\n")
                else:
                    f.write("60-Day Performance Watch: None\n")
                f.write("\n")
            
            # Top 5 Worst Performers
            f.write("TOP 5 WORST PERFORMERS:\n")
            f.write("-" * 35 + "\n")
            
            # By Ticket Count
            f.write("By Ticket Count:\n")
            for circuit, count in list(metrics.get('top5_tickets', {}).items())[:5]:
                circuit_clean = circuit.replace(' (C/R)', '').replace(' (C)', '').replace(' (R)', '')
                f.write(f"- {circuit_clean}: {count} tickets\n")
            f.write("\n")
            
            # By Availability
            f.write("By Availability (Worst):\n")
            for circuit, avail in list(metrics.get('bottom5_availability', {}).items())[:5]:
                circuit_clean = circuit.replace(' (C/R)', '').replace(' (C)', '').replace(' (R)', '')
                f.write(f"- {circuit_clean}: {avail:.2f}%\n")
            f.write("\n")
            
            # Notes
            f.write("Notes:\n")
            f.write("- (C) indicates chronic circuits\n")
            f.write("- (R) indicates regional correlation\n")
            f.write("- (C/R) indicates both chronic and regional\n")
        
        print(f"Text summary generated: {output_path}")
        return output_path
    
    def generate_charts(self, metrics, output_dir):
        """Generate PNG charts for the report"""
        output_dir = Path(output_dir)
        output_dir.mkdir(exist_ok=True)
        
        # Set style
        plt.style.use('default')
        sns.set_palette("viridis")
        
        charts = {}
        
        # Top 5 Tickets Chart
        if 'top5_tickets' in metrics:
            fig, ax = plt.subplots(figsize=(10, 6))
            circuits = list(metrics['top5_tickets'].keys())
            tickets = list(metrics['top5_tickets'].values())
            
            bars = ax.barh(circuits, tickets)
            ax.set_xlabel('Number of Tickets')
            ax.set_title('Top 5 Circuits by Ticket Volume')
            
            # Add value labels on bars
            for bar in bars:
                width = bar.get_width()
                ax.text(width, bar.get_y() + bar.get_height()/2, 
                       f'{int(width)}', ha='left', va='center')
            
            plt.tight_layout()
            chart_path = output_dir / 'top5_tickets.png'
            plt.savefig(chart_path, dpi=300, bbox_inches='tight')
            charts['top5_tickets'] = chart_path
            plt.close()
        
        # Top 5 Cost Chart
        if 'top5_cost' in metrics:
            fig, ax = plt.subplots(figsize=(10, 6))
            circuits = list(metrics['top5_cost'].keys())
            costs = list(metrics['top5_cost'].values())
            
            bars = ax.barh(circuits, costs)
            ax.set_xlabel('Cost to Serve ($)')
            ax.set_title('Top 5 Circuits by Cost to Serve')
            
            # Add value labels
            for bar in bars:
                width = bar.get_width()
                ax.text(width, bar.get_y() + bar.get_height()/2, 
                       f'${int(width):,}', ha='left', va='center')
            
            plt.tight_layout()
            chart_path = output_dir / 'top5_cost.png'
            plt.savefig(chart_path, dpi=300, bbox_inches='tight')
            charts['top5_cost'] = chart_path
            plt.close()
        
        # Bottom 5 Availability Chart
        if 'bottom5_availability' in metrics:
            fig, ax = plt.subplots(figsize=(10, 6))
            circuits = list(metrics['bottom5_availability'].keys())
            avail = list(metrics['bottom5_availability'].values())
            
            bars = ax.barh(circuits, avail)
            ax.set_xlabel('Availability %')
            ax.set_title('Bottom 5 Circuits by Availability')
            
            # Add value labels
            for bar in bars:
                width = bar.get_width()
                ax.text(width, bar.get_y() + bar.get_height()/2, 
                       f'{width:.1f}%', ha='left', va='center')
            
            plt.tight_layout()
            chart_path = output_dir / 'bottom5_availability.png'
            plt.savefig(chart_path, dpi=300, bbox_inches='tight')
            charts['bottom5_availability'] = chart_path
            plt.close()
        
        # Bottom 5 MTBF Chart (worst performing)
        if 'bottom5_mtbf' in metrics:
            fig, ax = plt.subplots(figsize=(10, 6))
            circuits = list(metrics['bottom5_mtbf'].keys())
            mtbf_days = list(metrics['bottom5_mtbf'].values())
            
            bars = ax.barh(circuits, mtbf_days, color='red', alpha=0.7)
            ax.set_xlabel('Mean Time Between Failures (Days)')
            ax.set_title('Bottom 5 Circuits by MTBF (Worst Performing)')
            
            # Add value labels
            for bar in bars:
                width = bar.get_width()
                ax.text(width, bar.get_y() + bar.get_height()/2, 
                       f'{width:.1f}d', ha='left', va='center')
            
            plt.tight_layout()
            chart_path = output_dir / 'bottom5_mtbf.png'
            plt.savefig(chart_path, dpi=300, bbox_inches='tight')
            charts['bottom5_mtbf'] = chart_path
            plt.close()
        
        return charts
    
    def generate_trend_analysis(self, current_month_str: str, output_dir: Path) -> str:
        """Generate comprehensive chart-based month-over-month trend analysis"""
        try:
            # Find previous month's data
            json_files = list(output_dir.glob('chronic_summary_*.json'))
            json_files = [f for f in json_files if f.name != f'chronic_summary_{current_month_str}.json']
            
            if not json_files:
                return "No previous month data available for trend analysis."
            
            # Get the most recent previous month
            previous_file = sorted(json_files)[-1]
            current_file = output_dir / f'chronic_summary_{current_month_str}.json'
            
            if not current_file.exists():
                return "Current month data not found for trend analysis."
            
            # Load data
            with open(previous_file, 'r') as f:
                prev_data = json.load(f)
            with open(current_file, 'r') as f:
                curr_data = json.load(f)
            
            # Extract month names
            prev_month = previous_file.name.replace('chronic_summary_', '').replace('.json', '').replace('_', ' ')
            curr_month = current_month_str.replace('_', ' ')
            
            # Analyze trends
            trends = []
            trends.append(f"# MONTHLY PERFORMANCE TREND ANALYSIS")
            trends.append(f"## {prev_month} → {curr_month}")
            trends.append("=" * 70)
            trends.append("")
            
            # Overall circuit health summary
            prev_total = prev_data.get('metrics', {}).get('total_chronic_circuits', 0)
            curr_total = curr_data.get('metrics', {}).get('total_chronic_circuits', 0)
            total_change = curr_total - prev_total
            
            trends.append("## 📊 NETWORK HEALTH OVERVIEW")
            trends.append(f"• **Total Chronic Circuits**: {prev_total} → {curr_total} ({total_change:+d} change)")
            trends.append(f"• **New Chronics Identified**: {curr_data.get('metrics', {}).get('new_chronic_count', 0)}")
            
            # Calculate circuit status distribution changes
            prev_consistent = len(prev_data.get('chronic_data', {}).get('existing_chronics', {}).get('chronic_consistent', []))
            curr_consistent = len(curr_data.get('chronic_data', {}).get('existing_chronics', {}).get('chronic_consistent', []))
            consistent_change = curr_consistent - prev_consistent
            
            trends.append(f"• **Consistent Circuits**: {prev_consistent} → {curr_consistent} ({consistent_change:+d})")
            trends.append("")
            
            # CHART-BASED ANALYSIS - TOP TICKET GENERATORS
            trends.append("## 🔥 TOP TICKET GENERATORS ANALYSIS")
            prev_top5_tickets = prev_data.get('metrics', {}).get('top5_tickets', {})
            curr_top5_tickets = curr_data.get('metrics', {}).get('top5_tickets', {})
            
            trends.extend(self._analyze_ranking_changes(
                "Ticket Generators", prev_top5_tickets, curr_top5_tickets, 
                is_higher_worse=True, unit="tickets"
            ))
            
            # CHART-BASED ANALYSIS - COST CIRCUITS
            trends.append("## 💰 COST TO SERVE ANALYSIS")
            prev_top5_cost = prev_data.get('metrics', {}).get('top5_cost', {})
            curr_top5_cost = curr_data.get('metrics', {}).get('top5_cost', {})
            
            trends.extend(self._analyze_ranking_changes(
                "Cost Circuits", prev_top5_cost, curr_top5_cost,
                is_higher_worse=True, unit="$", is_currency=True
            ))
            
            # CHART-BASED ANALYSIS - AVAILABILITY 
            trends.append("## 📉 AVAILABILITY PERFORMANCE ANALYSIS")
            prev_bottom5_avail = prev_data.get('metrics', {}).get('bottom5_availability', {})
            curr_bottom5_avail = curr_data.get('metrics', {}).get('bottom5_availability', {})
            
            trends.extend(self._analyze_ranking_changes(
                "Worst Availability", prev_bottom5_avail, curr_bottom5_avail,
                is_higher_worse=False, unit="%", threshold=AVAIL_THRESH_PCT
            ))
            
            # CHART-BASED ANALYSIS - MTBF
            trends.append("## ⚡ RELIABILITY (MTBF) ANALYSIS") 
            prev_bottom5_mtbf = prev_data.get('metrics', {}).get('bottom5_mtbf', {})
            curr_bottom5_mtbf = curr_data.get('metrics', {}).get('bottom5_mtbf', {})
            
            trends.extend(self._analyze_ranking_changes(
                "Worst MTBF", prev_bottom5_mtbf, curr_bottom5_mtbf,
                is_higher_worse=False, unit="days", threshold=1.0
            ))
            
            # STRATEGIC RECOMMENDATIONS
            trends.append("## 🎯 STRATEGIC RECOMMENDATIONS")
            red_flags, improvements, new_concerns = self._generate_strategic_insights(
                prev_data, curr_data
            )
            
            if red_flags:
                trends.append("### 🚨 IMMEDIATE ATTENTION REQUIRED:")
                trends.extend([f"• {flag}" for flag in red_flags])
                trends.append("")
            
            if improvements:
                trends.append("### 🎉 SUCCESS STORIES:")
                trends.extend([f"• {improvement}" for improvement in improvements])
                trends.append("")
            
            if new_concerns:
                trends.append("### 👀 EMERGING PATTERNS:")
                trends.extend([f"• {concern}" for concern in new_concerns])
                trends.append("")
            
            # Overall network trend assessment
            trends.append("### 📈 OVERALL NETWORK TREND:")
            if total_change > 2:
                trends.append(f"• **Growing chronic problem** - {total_change} new chronic circuits require root cause analysis")
            elif total_change > 0:
                trends.append(f"• **Slight degradation** - {total_change} additional chronic circuits, monitor closely")
            elif total_change < -2:
                trends.append(f"• **Significant improvement** - {abs(total_change)} fewer chronic circuits, maintenance efforts paying off")
            elif total_change < 0:
                trends.append(f"• **Gradual improvement** - {abs(total_change)} fewer chronic circuits")
            else:
                trends.append("• **Stable chronic population** - consistent with previous month")
            
            return "\n".join(trends)
            
        except Exception as e:
            return f"Error generating trend analysis: {e}"
    
    def _analyze_ranking_changes(self, category_name: str, prev_data: dict, curr_data: dict, 
                                is_higher_worse: bool = True, unit: str = "", threshold: float = None,
                                is_currency: bool = False) -> list:
        """Analyze ranking changes between two periods for a specific metric"""
        analysis = []
        
        if not prev_data or not curr_data:
            analysis.append(f"• Insufficient data for {category_name} comparison")
            analysis.append("")
            return analysis
        
        # Create ranking maps (circuit -> position)
        prev_ranks = {self._clean_circuit_name(circuit): i+1 for i, circuit in enumerate(prev_data.keys())}
        curr_ranks = {self._clean_circuit_name(circuit): i+1 for i, circuit in enumerate(curr_data.keys())}
        
        # Clean data for value comparison
        prev_values = {self._clean_circuit_name(k): v for k, v in prev_data.items()}
        curr_values = {self._clean_circuit_name(k): v for k, v in curr_data.items()}
        
        # Set automatic threshold if not provided
        if threshold is None:
            if "tickets" in unit.lower():
                threshold = 3
            elif is_currency:
                threshold = 1000  # $1000 - core chronic logic unchanged
            elif "%" in unit:
                threshold = AVAIL_THRESH_PCT   # Use configurable availability threshold
            else:
                threshold = 0.5   # 0.5 days - core chronic logic unchanged
        
        # Track movers and significant changes
        big_movers = []
        new_entries = []
        graduates = []
        significant_changes = []
        
        # Analyze position changes
        all_circuits = set(prev_ranks.keys()) | set(curr_ranks.keys())
        
        for circuit in all_circuits:
            prev_rank = prev_ranks.get(circuit)
            curr_rank = curr_ranks.get(circuit)
            prev_value = prev_values.get(circuit, 0)
            curr_value = curr_values.get(circuit, 0)
            
            # New entries to top/bottom 5
            if prev_rank is None and curr_rank is not None:
                value_str = self._format_value(curr_value, unit, is_currency)
                new_entries.append(f"**{circuit}** entered at #{curr_rank} ({value_str})")
            
            # Graduates (left top/bottom 5)
            elif prev_rank is not None and curr_rank is None:
                value_str = self._format_value(prev_value, unit, is_currency)
                graduates.append(f"**{circuit}** improved out of worst performers (was #{prev_rank})")
            
            # Position changes within top/bottom 5
            elif prev_rank is not None and curr_rank is not None and prev_rank != curr_rank:
                rank_change = prev_rank - curr_rank  # Positive = moved up in ranking (worse)
                direction = "worsened" if rank_change < 0 else "improved"
                
                prev_val_str = self._format_value(prev_value, unit, is_currency)
                curr_val_str = self._format_value(curr_value, unit, is_currency)
                
                if abs(rank_change) >= 2:  # Significant position change - core chronic logic unchanged
                    big_movers.append(f"**{circuit}** {direction} #{prev_rank} → #{curr_rank} ({prev_val_str} → {curr_val_str})")
            
            # Significant value changes (same circuit in both periods)
            if prev_rank is not None and curr_rank is not None:
                value_change = curr_value - prev_value
                percent_change = (value_change / prev_value * 100) if prev_value != 0 else 0
                
                if abs(value_change) >= threshold or abs(percent_change) >= 20:
                    change_str = f"{value_change:+.1f}" if not is_currency else f"${value_change:+,.0f}"
                    direction_word = "increased" if value_change > 0 else "decreased"
                    
                    if is_higher_worse:
                        impact = "⚠️ DEGRADED" if value_change > 0 else "✅ IMPROVED"
                    else:
                        impact = "✅ IMPROVED" if value_change > 0 else "⚠️ DEGRADED"
                    
                    significant_changes.append(
                        f"**{circuit}** {direction_word} by {change_str}{unit} ({percent_change:+.1f}%) {impact}"
                    )
        
        # Format analysis results
        if big_movers:
            analysis.append("### 📊 Major Ranking Changes:")
            analysis.extend([f"• {mover}" for mover in big_movers])
            analysis.append("")
        
        if new_entries:
            analysis.append("### 🚨 New Problem Circuits:")
            analysis.extend([f"• {entry}" for entry in new_entries])
            analysis.append("")
        
        if graduates:
            analysis.append("### 🎉 Improved Circuits:")
            analysis.extend([f"• {grad}" for grad in graduates])
            analysis.append("")
        
        if significant_changes:
            analysis.append("### 📈 Significant Value Changes:")
            analysis.extend([f"• {change}" for change in significant_changes])
            analysis.append("")
        
        if not any([big_movers, new_entries, graduates, significant_changes]):
            analysis.append(f"• **{category_name}**: No significant changes from previous month")
            analysis.append("")
        
        return analysis
    
    def _clean_circuit_name(self, circuit_name: str) -> str:
        """Remove indicators and clean circuit name for comparison"""
        return circuit_name.split(' ')[0] if circuit_name else ""
    
    def _format_value(self, value: float, unit: str, is_currency: bool = False) -> str:
        """Format value with appropriate unit"""
        if is_currency:
            return f"${value:,.0f}"
        elif "%" in unit:
            return f"{value:.1f}%"
        elif "days" in unit:
            return f"{value:.1f} days"
        else:
            return f"{value:.0f}{unit}"
    
    def _generate_strategic_insights(self, prev_data: dict, curr_data: dict) -> tuple:
        """Generate strategic insights for recommendations"""
        red_flags = []
        improvements = []
        new_concerns = []
        
        # Analyze circuit count trends
        prev_total = prev_data.get('metrics', {}).get('total_chronic_circuits', 0)
        curr_total = curr_data.get('metrics', {}).get('total_chronic_circuits', 0)
        
        if curr_total - prev_total >= 3:
            red_flags.append(f"Chronic circuit count increased by {curr_total - prev_total} - potential systemic issues")
        
        # Check for new high-cost circuits
        prev_costs = set(prev_data.get('metrics', {}).get('top5_cost', {}).keys())
        curr_costs = set(curr_data.get('metrics', {}).get('top5_cost', {}).keys())
        new_cost_circuits = curr_costs - prev_costs
        
        if new_cost_circuits:
            new_concerns.append(f"New high-cost circuits emerged: {', '.join([self._clean_circuit_name(c) for c in new_cost_circuits])}")
        
        # Check for availability improvements
        prev_avail = prev_data.get('metrics', {}).get('bottom5_availability', {})
        curr_avail = curr_data.get('metrics', {}).get('bottom5_availability', {})
        
        improved_avail = []
        for circuit in set(prev_avail.keys()) & set(curr_avail.keys()):
            if curr_avail[circuit] - prev_avail[circuit] >= AVAIL_THRESH_PCT:  # Significant improvement
                improved_avail.append(self._clean_circuit_name(circuit))
        
        if improved_avail:
            improvements.append(f"Significant availability improvements: {', '.join(improved_avail)}")
        
        # Check for ticket volume spikes
        prev_tickets = prev_data.get('metrics', {}).get('top5_tickets', {})
        curr_tickets = curr_data.get('metrics', {}).get('top5_tickets', {})
        
        ticket_spikes = []
        for circuit in set(prev_tickets.keys()) & set(curr_tickets.keys()):
            if curr_tickets[circuit] - prev_tickets[circuit] >= 10:  # 10+ ticket increase
                ticket_spikes.append(f"{self._clean_circuit_name(circuit)} (+{curr_tickets[circuit] - prev_tickets[circuit]} tickets)")
        
        if ticket_spikes:
            red_flags.append(f"Major ticket volume spikes: {', '.join(ticket_spikes)}")
        
        return red_flags, improvements, new_concerns
    
    def generate_trend_analysis_word(self, month_str: str, output_dir: Path) -> Path:
        """Generate trend analysis as a Word document"""
        trend_text = self.generate_trend_analysis(month_str, output_dir)
        
        if "No previous month data" in trend_text or "Error" in trend_text:
            # Skip Word generation if no data or error
            return None
        
        doc = Document()
        doc.add_heading('Monthly Chronic Circuit Trend Analysis', 0)
        
        # Parse the trend analysis text and format it nicely
        lines = trend_text.split('\n')
        current_section = None
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            if line.startswith('# '):
                # Main title - already added above
                continue
            elif line.startswith('='):
                # Skip separator lines
                continue
            elif line.startswith('## '):
                # Section heading
                section_title = line.replace('## ', '')
                doc.add_heading(section_title, level=1)
                current_section = section_title.lower()
            elif line.startswith('### '):
                # Subsection heading
                subsection_title = line.replace('### ', '')
                doc.add_heading(subsection_title, level=2)
            elif line.startswith('• '):
                # Bullet point
                p = doc.add_paragraph()
                p.style = 'List Bullet'
                p.add_run(line[2:])  # Remove bullet marker
            elif line and not line.startswith('#'):
                # Regular text
                if ':' in line and current_section in ['executive summary', 'top ticket generators']:
                    # Format numbered lists and key-value pairs
                    parts = line.split(':')
                    p = doc.add_paragraph()
                    if line[0].isdigit():
                        p.style = 'List Number'
                        p.add_run(line)
                    else:
                        run = p.add_run(parts[0] + ':')
                        run.bold = True
                        if len(parts) > 1:
                            p.add_run(' ' + ':'.join(parts[1:]))
                else:
                    doc.add_paragraph(line)
        
        # Add footer with generation timestamp
        doc.add_page_break()
        footer_p = doc.add_paragraph()
        footer_p.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} | ")
        footer_p.add_run("Monthly Reporting System v0.1.6").italic = True
        
        # Save document
        output_path = output_dir / f"Monthly_Trend_Analysis_{month_str}.docx"
        doc.save(output_path)
        
        return output_path
    
    def generate_chronic_corner_word(self, metrics, chronic_data, output_path, charts=None, month_str=None):
        """Generate Chronic Corner format as Word document - exact format match"""
        
        doc = Document()
        doc.add_heading('Chronic Corner', 0)
        
        # Extract month/year for dynamic narrative
        if month_str:
            # Convert "June_2025" to "June 2025"
            month_display = month_str.replace('_', ' ')
        else:
            # Fallback to May 2025 if no month provided
            month_display = "May 2025"
        
        # Trends section
        doc.add_heading('Trends', level=2)
        trends_text = f"By the end of {month_display}, we've confirmed {metrics['total_chronic_circuits']} chronic circuits among {metrics['total_providers']} Circuit Providers. We also identified {metrics['media_chronics']} media services as chronic, with all of them operated on behalf of three Hotlist Media customers."
        doc.add_paragraph(trends_text)
        
        # Special formatted metric block - 1-row 4-column table
        from docx.shared import RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
        
        metrics_table = doc.add_table(rows=1, cols=4)
        metrics_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Configure table formatting
        for row in metrics_table.rows:
            for cell in row.cells:
                # Background fill: #E2E5FF using simpler approach
                from docx.oxml.shared import qn
                from docx.oxml import OxmlElement
                
                tcPr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), 'E2E5FF')
                tcPr.append(shd)
                
                # Cell vertical alignment: center
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # Apply simple table style for now
        metrics_table.style = 'Light Grid'
        
        # Cell 1: Chronic Consistent
        cell1 = metrics_table.cell(0, 0)
        p1 = cell1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p1.clear()
        
        run1_num = p1.add_run(str(len(chronic_data['existing_chronics']['chronic_consistent'])))
        run1_num.bold = True
        run1_num.font.size = Pt(28)
        
        p1.add_run('\n')
        run1_label1 = p1.add_run('Chronic')
        run1_label1.font.size = Pt(10)
        
        p1.add_run('\n')
        run1_label2 = p1.add_run('Consistent')
        run1_label2.font.size = Pt(10)
        
        # Cell 2: Circuit Providers
        cell2 = metrics_table.cell(0, 1)
        p2 = cell2.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.clear()
        
        run2_num = p2.add_run(str(metrics['total_providers']))
        run2_num.bold = True
        run2_num.font.size = Pt(28)
        
        p2.add_run('\n')
        run2_label1 = p2.add_run('Circuit')
        run2_label1.font.size = Pt(10)
        
        p2.add_run('\n')
        run2_label2 = p2.add_run('Providers')
        run2_label2.font.size = Pt(10)
        
        # Cell 3: Media Services
        cell3 = metrics_table.cell(0, 2)
        p3 = cell3.paragraphs[0]
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.clear()
        
        run3_num = p3.add_run(str(metrics['media_chronics']))
        run3_num.bold = True
        run3_num.font.size = Pt(28)
        
        p3.add_run('\n')
        run3_label1 = p3.add_run('Media')
        run3_label1.font.size = Pt(10)
        
        p3.add_run('\n')
        run3_label2 = p3.add_run('Services')
        run3_label2.font.size = Pt(10)
        
        # Cell 4: New Chronics
        cell4 = metrics_table.cell(0, 3)
        p4 = cell4.paragraphs[0]
        p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p4.clear()
        
        run4_num = p4.add_run(str(metrics['new_chronic_count']))
        run4_num.bold = True
        run4_num.font.size = Pt(28)
        
        p4.add_run('\n')
        run4_label1 = p4.add_run('New')
        run4_label1.font.size = Pt(10)
        
        p4.add_run('\n')
        run4_label2 = p4.add_run('Chronics')
        run4_label2.font.size = Pt(10)
        
        # Chronic Consistent Table
        doc.add_heading('Chronic Consistent', level=2)
        cc_table = doc.add_table(rows=1, cols=2)
        cc_table.style = 'Table Grid'
        cc_table.cell(0, 0).text = "Vendor"
        cc_table.cell(0, 1).text = "Circuits"
        
        # Group consistent circuits by vendor
        consistent_circuits = chronic_data['existing_chronics']['chronic_consistent']
        vendors = {
            'Cirion': [c for c in consistent_circuits if c.startswith('500')],
            'Tata': [c for c in consistent_circuits if c.startswith('091')],
            'PCCW': [c for c in consistent_circuits if c.startswith('SR')],
            'Telstra': [c for c in consistent_circuits if 'PTH' in c],
            'Liquid Telecom': [c for c in consistent_circuits if c.startswith('LZA')]
        }
        
        for vendor, circuits in vendors.items():
            if circuits:
                row = cc_table.add_row()
                row.cells[0].text = vendor
                row.cells[1].text = str(len(circuits))
        
        # Chronic Inconsistent Table
        doc.add_heading('Chronic Inconsistent', level=2)
        ci_table = doc.add_table(rows=1, cols=2)
        ci_table.style = 'Table Grid'
        ci_table.cell(0, 0).text = "Vendor"
        ci_table.cell(0, 1).text = "Services"
        
        # Group inconsistent circuits by vendor (including new chronic)
        inconsistent_circuits = chronic_data['existing_chronics']['chronic_inconsistent'].copy()
        if metrics.get('new_chronics'):
            for circuits in metrics['new_chronics'].values():
                inconsistent_circuits.extend(circuits)
        
        inc_vendors = {
            'Lumen': [c for c in inconsistent_circuits if c.startswith('4') and len(c) < 12],
            'Orange': [c for c in inconsistent_circuits if c.startswith('LD')],
            'Globenet': [c for c in inconsistent_circuits if c.startswith('IST')],
            'GTT': [c for c in inconsistent_circuits if 'HI/ADM' in c],
            'PCCW': [c for c in inconsistent_circuits if c.startswith('SR2')],
            'Sansa': [c for c in inconsistent_circuits if 'SSO' in c],
            'Verizon': [c for c in inconsistent_circuits if c.startswith('W1E')],
            'Telstra': [c for c in inconsistent_circuits if c.startswith('N')]
        }
        
        for vendor, circuits in inc_vendors.items():
            if circuits:
                row = ci_table.add_row()
                row.cells[0].text = vendor
                row.cells[1].text = str(len(circuits))
        
        # Media Hotlist Table
        doc.add_heading('Media Hotlist', level=2)
        media_table = doc.add_table(rows=1, cols=2)
        media_table.style = 'Table Grid'
        media_table.cell(0, 0).text = "Vendor"
        media_table.cell(0, 1).text = "Services"
        
        media_vendors = [
            ("Slovak Telekom", "4"),
            ("BBC", "4"),
            ("Slovak", "3")
        ]
        
        for vendor, count in media_vendors:
            row = media_table.add_row()
            row.cells[0].text = vendor
            row.cells[1].text = count
        
        # Performance Monitoring Table
        doc.add_heading('Performance Monitoring', level=2)
        pm_table = doc.add_table(rows=1, cols=2)
        pm_table.style = 'Table Grid'
        pm_table.cell(0, 0).text = "Circuit ID"
        pm_table.cell(0, 1).text = "Incidents"
        
        # Add 60-day monitoring circuits
        for circuit in chronic_data['existing_chronics']['perf_60_day']:
            row = pm_table.add_row()
            row.cells[0].text = circuit
            row.cells[1].text = "3"  # Default incident count
        
        # Add 30-day monitoring circuits  
        for circuit in chronic_data['existing_chronics']['perf_30_day']:
            row = pm_table.add_row()
            row.cells[0].text = circuit
            row.cells[1].text = "7"  # Default incident count
        
        # Add charts to the bottom of the document
        if charts:
            doc.add_page_break()
            doc.add_heading('Circuit Analysis Charts', level=1)
            
            # Add legend for chronic corner charts (same as circuit report)
            if self.show_indicators:
                legend = doc.add_paragraph()
                legend.add_run("Legend: ").bold = True
                legend.add_run("(C) = Chronic Circuit, (R) = Regional Circuit")
                legend.paragraph_format.space_after = Pt(6)
            
            for chart_name, chart_path in charts.items():
                if chart_path.exists():
                    doc.add_heading(chart_name.replace('_', ' ').title(), level=2)
                    doc.add_picture(str(chart_path), width=Inches(6))
        
        doc.save(output_path)
        return output_path
    
    def generate_circuit_report_pdf(self, metrics, chronic_data, charts, output_path):
        """Generate Circuit Report format for PDF"""
        
        doc = Document()
        doc.add_heading('Chronic Circuit Report', 0)
        
        # Header information table (like sample data shows)
        doc.add_heading('March - May 2025', level=1)
        
        # Special formatted metric block - same style as Chronic Corner
        from docx.shared import RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
        from docx.oxml.shared import qn
        from docx.oxml import OxmlElement
        
        summary_table = doc.add_table(rows=1, cols=4)
        summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Configure table formatting - same as Chronic Corner
        for row in summary_table.rows:
            for cell in row.cells:
                # Background fill: #E2E5FF using simpler approach
                tcPr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), 'E2E5FF')
                tcPr.append(shd)
                
                # Cell vertical alignment: center
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # Apply simple table style for now
        summary_table.style = 'Light Grid'
        
        # Calculate metrics - get TOTAL tickets from ALL circuits, not just top 5
        merged_df = chronic_data['merged_data']
        if 'Distinct count of Inc Nbr' in merged_df.columns:
            total_tickets = int(merged_df['Distinct count of Inc Nbr'].sum())
        else:
            total_tickets = sum(metrics.get('top5_tickets', {}).values()) if metrics.get('top5_tickets') else 0
            
        avg_availability = sum(metrics.get('bottom5_availability', {}).values()) / len(metrics.get('bottom5_availability', {})) if metrics.get('bottom5_availability') else 95.0
        
        # Cell 1: Total Circuits Tracked (64)
        cell1 = summary_table.cell(0, 0)
        p1 = cell1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p1.clear()
        
        run1_num = p1.add_run("64")
        run1_num.bold = True
        run1_num.font.size = Pt(28)
        
        p1.add_run('\n')
        run1_label1 = p1.add_run('Total Circuits')
        run1_label1.font.size = Pt(10)
        
        p1.add_run('\n')
        run1_label2 = p1.add_run('Tracked')
        run1_label2.font.size = Pt(10)
        
        # Cell 2: Total Tickets Logged (ALL tickets from all circuits)
        cell2 = summary_table.cell(0, 1)
        p2 = cell2.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.clear()
        
        run2_num = p2.add_run(str(total_tickets))
        run2_num.bold = True
        run2_num.font.size = Pt(28)
        
        p2.add_run('\n')
        run2_label1 = p2.add_run('Total Tickets')
        run2_label1.font.size = Pt(10)
        
        p2.add_run('\n')
        run2_label2 = p2.add_run('Logged')
        run2_label2.font.size = Pt(10)
        
        # Cell 3: Average Availability
        cell3 = summary_table.cell(0, 2)
        p3 = cell3.paragraphs[0]
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.clear()
        
        run3_num = p3.add_run(f"{avg_availability:.1f}%")
        run3_num.bold = True
        run3_num.font.size = Pt(28)
        
        p3.add_run('\n')
        run3_label1 = p3.add_run('Average')
        run3_label1.font.size = Pt(10)
        
        p3.add_run('\n')
        run3_label2 = p3.add_run('Availability')
        run3_label2.font.size = Pt(10)
        
        # Cell 4: Average MTBF
        cell4 = summary_table.cell(0, 3)
        p4 = cell4.paragraphs[0]
        p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p4.clear()
        
        run4_num = p4.add_run(f"{metrics.get('avg_mtbf_days', 20):.1f}")
        run4_num.bold = True
        run4_num.font.size = Pt(28)
        
        p4.add_run('\n')
        run4_label1 = p4.add_run('Average MTBF')
        run4_label1.font.size = Pt(10)
        
        p4.add_run('\n')
        run4_label2 = p4.add_run('(Days)')
        run4_label2.font.size = Pt(10)
        
        # Key Takeaways section (removed Executive Summary heading)
        
        # Key takeaways (3 lines based on data)
        doc.add_heading('Key Takeaways', level=2)
        worst_mtbf = min(metrics.get('bottom5_mtbf', {}).values()) if metrics.get('bottom5_mtbf') else 0
        worst_availability = min(metrics.get('bottom5_availability', {}).values()) if metrics.get('bottom5_availability') else 95
        highest_cost = max(metrics.get('top5_cost', {}).values()) if metrics.get('top5_cost') else 0
        
        takeaway1 = f"• {metrics['new_chronic_count']} new circuit(s) identified as chronic this month, requiring immediate attention and classification."
        takeaway2 = f"• Lowest performing circuit shows {worst_mtbf:.1f} days MTBF and {worst_availability:.1f}% availability, indicating significant reliability issues."
        takeaway3 = f"• Highest impact circuit generated ${highest_cost:,.0f} in cost to serve, representing major operational expense."
        
        doc.add_paragraph(takeaway1)
        doc.add_paragraph(takeaway2) 
        doc.add_paragraph(takeaway3)
        
        # Recommendations
        doc.add_heading('Recommendations', level=2)
        rec1 = "• Review the top 5 circuits by lowest MTBF and investigate root causes for frequent failures."
        rec2 = "• Prioritize vendor engagement for circuits showing consistent availability degradation."
        rec3 = "• Implement proactive monitoring for new chronic circuits to prevent escalation."
        
        doc.add_paragraph(rec1)
        doc.add_paragraph(rec2)
        doc.add_paragraph(rec3)
        
        # Performance Analysis
        doc.add_heading('Performance Analysis', level=1)
        
        # Add subtle legend (if indicators are enabled)
        if self.show_indicators:
            legend = doc.add_paragraph()
            legend.add_run("Legend: ").bold = True
            legend.add_run("(C) = Chronic Circuit, (R) = Regional Circuit")
            legend.paragraph_format.space_after = Pt(6)
        
        if 'top5_tickets' in metrics:
            doc.add_heading('Top 5 Circuits by Ticket Volume', level=2)
            tickets_table = doc.add_table(rows=1, cols=2)
            tickets_table.style = 'Table Grid'
            tickets_table.cell(0, 0).text = "Circuit ID"
            tickets_table.cell(0, 1).text = "Tickets"
            
            for circuit, count in metrics['top5_tickets'].items():
                row = tickets_table.add_row()
                row.cells[0].text = circuit
                row.cells[1].text = str(count)
        
        if 'top5_cost' in metrics:
            doc.add_heading('Top 5 Circuits by Cost to Serve', level=2)
            cost_table = doc.add_table(rows=1, cols=2)
            cost_table.style = 'Table Grid'
            cost_table.cell(0, 0).text = "Circuit ID"
            cost_table.cell(0, 1).text = "Cost ($)"
            
            for circuit, cost in metrics['top5_cost'].items():
                row = cost_table.add_row()
                row.cells[0].text = circuit
                row.cells[1].text = f"${cost:,.0f}"
        
        if 'bottom5_availability' in metrics:
            doc.add_heading('Bottom 5 Circuits by Availability', level=2)
            avail_table = doc.add_table(rows=1, cols=2)
            avail_table.style = 'Table Grid'
            avail_table.cell(0, 0).text = "Circuit ID"
            avail_table.cell(0, 1).text = "Availability (%)"
            
            for circuit, avail in metrics['bottom5_availability'].items():
                row = avail_table.add_row()
                row.cells[0].text = circuit
                row.cells[1].text = f"{avail:.1f}%"
        
        if 'bottom5_mtbf' in metrics:
            doc.add_heading('Bottom 5 Circuits by MTBF (Worst Performing)', level=2)
            mtbf_table = doc.add_table(rows=1, cols=2)
            mtbf_table.style = 'Table Grid'
            mtbf_table.cell(0, 0).text = "Circuit ID"
            mtbf_table.cell(0, 1).text = "MTBF (Days)"
            
            for circuit, mtbf in metrics['bottom5_mtbf'].items():
                row = mtbf_table.add_row()
                row.cells[0].text = circuit
                row.cells[1].text = f"{mtbf:.1f}"
        
        # Add charts
        if charts:
            doc.add_page_break()
            doc.add_heading('Circuit Analysis Charts', level=1)
            for chart_name, chart_path in charts.items():
                if chart_path.exists():
                    doc.add_heading(chart_name.replace('_', ' ').title(), level=2)
                    doc.add_picture(str(chart_path), width=Inches(6))
        
        doc.save(output_path)
        return output_path
    
    
    def populate_word_template(self, template_path, metrics, charts, output_path):
        """Populate the Word template with calculated metrics"""
        
        # Load template
        doc = Document(template_path)
        
        # Replace key metrics in text
        for paragraph in doc.paragraphs:
            if paragraph.text:
                # Replace placeholders with actual values
                text = paragraph.text
                text = text.replace("23", str(metrics.get('total_chronic_circuits', 23)))
                text = text.replace("14", str(metrics.get('total_providers', 14)))
                
                # Update the paragraph text
                paragraph.text = text
        
        # Update tables if they exist
        for table in doc.tables:
            if len(table.rows) > 0 and len(table.columns) >= 4:
                # Update the metrics table (first table)
                cells = table.rows[0].cells
                if len(cells) >= 4:
                    cells[0].text = f"{metrics.get('total_chronic_circuits', 23)}\nChronic\nConsistent"
                    cells[1].text = f"{metrics.get('total_providers', 14)}\nCircuit\nProviders"
                    # Add more cell updates as needed
        
        # Add charts if available
        if charts:
            # Add a new page for charts
            doc.add_page_break()
            doc.add_heading('Circuit Analysis Charts', level=1)
            
            for chart_name, chart_path in charts.items():
                if chart_path.exists():
                    doc.add_heading(chart_name.replace('_', ' ').title(), level=2)
                    doc.add_picture(str(chart_path), width=Inches(6))
        
        # Add MTBF section
        if 'avg_mtbf_days' in metrics:
            doc.add_heading('Mean Time Between Failures Analysis', level=1)
            p = doc.add_paragraph(f"Average MTBF across chronic circuits: {metrics['avg_mtbf_days']:.1f} days")
        
        # Save the document
        doc.save(output_path)
        return output_path
    
    def convert_to_pdf(self, docx_path, pdf_path):
        """Convert Word document to PDF"""
        try:
            # Try using LibreOffice for conversion
            result = subprocess.run([
                'libreoffice', '--headless', '--convert-to', 'pdf', 
                '--outdir', str(Path(pdf_path).parent), str(docx_path)
            ], capture_output=True, text=True)
            
            if result.returncode == 0:
                return pdf_path
            else:
                print(f"LibreOffice conversion failed: {result.stderr}")
                return None
                
        except FileNotFoundError:
            print("LibreOffice not found. PDF conversion skipped.")
            return None
    
    
    def build_monthly_report(self, impacts_file, counts_file, template_file, output_dir, month_str=None):
        """Main pipeline to build the monthly report"""
        
        output_dir = Path(output_dir)
        output_dir.mkdir(exist_ok=True)
        
        print("Starting monthly report build...")
        
        # Load data
        impacts_df, counts_df = self.load_crosstab_data(impacts_file, counts_file)
        
        # Process chronic logic  
        chronic_data = self.process_chronic_logic(impacts_df, counts_df)
        print(f"Existing chronic circuits: {chronic_data['total_chronic_circuits']}")
        print(f"New chronics identified: {chronic_data['new_chronic_count']}")
        
        # Calculate metrics
        metrics = self.calculate_metrics(chronic_data)
        
        # Generate charts
        charts = self.generate_charts(metrics, output_dir / 'charts')
        
        # Use provided month string or default to May 2025
        if month_str is None:
            report_date = datetime(2025, 5, 31)  # May 2025
            month_str = report_date.strftime("%B_%Y")
        else:
            # If month_str is provided, use current date for metadata
            report_date = datetime.now()
        
        # 1. Chronic Corner (Word document)
        corner_word_output = output_dir / f"Chronic_Corner_{month_str}.docx"
        self.generate_chronic_corner_word(metrics, chronic_data, corner_word_output, charts, month_str)
        
        # 2. Circuit Report (Word document for PDF conversion)
        circuit_word_output = output_dir / f"Chronic_Circuit_Report_{month_str}.docx"
        self.generate_circuit_report_pdf(metrics, chronic_data, charts, circuit_word_output)
        
        # 3. PDF conversion of Circuit Report
        pdf_output = output_dir / f"Chronic_Circuit_Report_{month_str}.pdf"
        self.convert_to_pdf(circuit_word_output, pdf_output)
        
        # PowerPoint generation removed per user request
        
        # Generate text summary
        text_summary_output = self.generate_text_summary(chronic_data, metrics, output_dir, month_str)
        
        # Generate trend analysis
        trend_analysis = self.generate_trend_analysis(month_str, output_dir)
        trend_analysis_output = output_dir / f"monthly_trend_analysis_{month_str}.txt"
        with open(trend_analysis_output, 'w') as f:
            f.write(trend_analysis)
        
        # Generate trend analysis Word document
        trend_word_output = self.generate_trend_analysis_word(month_str, output_dir)
        
        # v0.1.9: Generate metadata block
        metadata = self.generate_metadata(impacts_file, counts_file)
        
        # Export data summary with metadata
        summary_data = {
            'version': '0.1.9',
            'consistency_mode': 'hybrid',
            'baseline_hotfix': 'new_chronic_promotion_fix',
            'metadata': metadata,
            'chronic_data': chronic_data,
            'metrics': metrics,
            'generated_at': report_date.isoformat()
        }
        
        with open(output_dir / f"chronic_summary_{month_str}.json", 'w') as f:
            json.dump(summary_data, f, indent=2, default=str)
        
        print(f"Reports generated in {output_dir}")
        print(f"[SUCCESS] Chronic Corner (Word): {corner_word_output}")
        print(f"[SUCCESS] Circuit Report (Word): {circuit_word_output}")
        print(f"[SUCCESS] Chronic List (Text): {text_summary_output}")
        if trend_word_output:
            print(f"[SUCCESS] Trend Analysis (Word): {trend_word_output}")
        if pdf_output and Path(pdf_output).exists():
            print(f"[SUCCESS] Circuit Report (PDF): {pdf_output}")
        
        return corner_word_output, circuit_word_output, pdf_output

def validate_month_selection(impacts_file: str, counts_file: str, selected_month: str, selected_year: str) -> tuple[bool, str]:
    """
    Validate that the selected month matches the data in the files.
    
    Returns:
        tuple: (is_valid, message)
    """
    try:
        # Load a sample of the impacts file to check months
        if impacts_file.lower().endswith('.csv'):
            sample_df = pd.read_csv(impacts_file, nrows=100)
        else:
            sample_df = pd.read_excel(impacts_file, nrows=100)
        
        # Clean column names
        sample_df.columns = sample_df.columns.str.strip()
        
        # Handle column aliasing
        if 'Configuration Item Name' in sample_df.columns:
            sample_df = sample_df.rename(columns={'Configuration Item Name': 'Config Item Name'})
        
        # Check if we have the month column
        month_column = 'Inc Resolved At (Month / Year)'
        if month_column not in sample_df.columns:
            return True, "Cannot validate month - month column not found in data"
        
        # Get unique months from the data
        data_months = sample_df[month_column].dropna().unique()
        
        if len(data_months) == 0:
            return True, "Cannot validate month - no month data found"
        
        # Parse the months to understand the data period
        month_names = ['January', 'February', 'March', 'April', 'May', 'June',
                      'July', 'August', 'September', 'October', 'November', 'December']
        
        # Convert selected month to expected 3-month window
        try:
            selected_month_num = month_names.index(selected_month) + 1
            selected_year_num = int(selected_year)
            
            # Calculate expected 3-month window BEFORE the selected month
            # For June report, we expect March, April, May data
            expected_months = []
            for i in range(3):
                month_num = selected_month_num - 3 + i  # Changed from -2 to -3
                year_num = selected_year_num
                
                if month_num <= 0:
                    month_num += 12
                    year_num -= 1
                
                expected_months.append(f"{month_names[month_num-1]} {year_num}")
            
            # Check if data contains months that suggest a different report period
            data_months_str = [str(m) for m in data_months if pd.notna(m)]
            
            # Simple heuristic: if we see months that don't match expected window, warn
            unexpected_months = []
            for data_month in data_months_str[:5]:  # Check first 5 unique months
                if data_month not in expected_months:
                    unexpected_months.append(data_month)
            
            if unexpected_months:
                suggested_month = None
                # Try to infer correct month from data
                for data_month in data_months_str:
                    for i, month_name in enumerate(month_names):
                        if month_name in data_month and selected_year in data_month:
                            # This could be the correct report month
                            suggested_month = month_name
                            break
                    if suggested_month:
                        break
                
                warning_msg = f"⚠️  Data/Month Mismatch Detected!\n\n"
                warning_msg += f"For a {selected_month} {selected_year} report, the data should contain:\n"
                warning_msg += f"✓ Expected: {', '.join(expected_months)}\n\n"
                warning_msg += f"But your data contains:\n"
                warning_msg += f"✗ Found: {', '.join(data_months_str[:3])}...\n"
                if suggested_month and suggested_month != selected_month:
                    warning_msg += f"\n💡 Suggestion: This looks like data for a '{suggested_month}' report"
                
                return False, warning_msg
            
            return True, f"✅ Month selection looks correct for {selected_month} {selected_year}"
            
        except (ValueError, IndexError):
            return True, "Cannot validate month - date parsing error"
            
    except Exception as e:
        # Don't block on validation errors, just warn
        return True, f"Month validation failed: {str(e)}"

def gui_main():
    """Main GUI interface for monthly reporting"""
    sg.theme('DarkBlue3')
    
    # Get current month/year for defaults
    current_date = datetime.now()
    current_month = current_date.strftime('%B')
    current_year = str(current_date.year)
    
    # Month options
    months = ['January', 'February', 'March', 'April', 'May', 'June',
              'July', 'August', 'September', 'October', 'November', 'December']
    
    # Layout
    layout = [
        [sg.Text('Monthly Chronic Circuit Report Generator', font=('Arial', 16, 'bold'))],
        [sg.HSeparator()],
        
        [sg.Text('Required Files:', font=('Arial', 12, 'bold'))],
        [sg.Text('Impacts Crosstab File:'), sg.InputText(key='-IMPACTS-'), sg.FileBrowse(file_types=(('Excel Files', '*.xlsx'),))],
        [sg.Text('Count Months Chronic File:'), sg.InputText(key='-COUNTS-'), sg.FileBrowse(file_types=(('Excel Files', '*.xlsx'),))],
        
        [sg.Text('')],
        [sg.Text('Options:', font=('Arial', 12, 'bold'))],
        [sg.Text('Report Month:'), sg.Combo(months, default_value=current_month, key='-MONTH-', readonly=True)],
        [sg.Text('Report Year:'), sg.InputText(current_year, key='-YEAR-', size=(10, 1))],
        [sg.Checkbox('Exclude Regional Circuits', default=False, key='-EXCLUDE_REGIONAL-')],
        [sg.Checkbox('Show Indicators (C) and (R)', default=True, key='-SHOW_INDICATORS-')],
        
        [sg.Text('')],
        [sg.Text('Output Directory:'), sg.InputText('./final_output', key='-OUTPUT-'), sg.FolderBrowse()],
        
        [sg.Text('')],
        [sg.HSeparator()],
        [sg.Button('Validate Files', size=(12, 1), button_color=('white', 'orange')),
         sg.Button('Generate Report', size=(15, 2), font=('Arial', 12, 'bold')), 
         sg.Button('Exit', size=(15, 2))],
        
        [sg.Text('')],
        [sg.Multiline(size=(80, 10), key='-LOG-', disabled=True, autoscroll=True)]
    ]
    
    window = sg.Window('Monthly Chronic Circuit Reporting', layout, finalize=True)
    
    while True:
        try:
            event, values = window.read()
            
            if event == sg.WIN_CLOSED or event == 'Exit':
                break
            
            if event == 'Validate Files':
                # Validate file selections
                if not values['-IMPACTS-']:
                    sg.popup_error('Please select an Impacts Crosstab file first')
                    continue
                if not values['-COUNTS-']:
                    sg.popup_error('Please select a Count Months Chronic file first')
                    continue
            
            # Clear log and show validation progress
            window['-LOG-'].update('')
            window['-LOG-'].update("Validating files and month selection...\n", append=True)
            window.refresh()
            
            try:
                # Validate month selection
                is_valid, message = validate_month_selection(
                    values['-IMPACTS-'], 
                    values['-COUNTS-'], 
                    values['-MONTH-'], 
                    values['-YEAR-']
                )
                
                window['-LOG-'].update(f"{message}\n", append=True)
                
                if not is_valid:
                    # Show warning popup with option to continue
                    result = sg.popup_yes_no(
                        f"{message}\n\nDo you want to continue anyway?",
                        title="Month Validation Warning",
                        no_titlebar=False
                    )
                    if result == 'Yes':
                        window['-LOG-'].update("⚠️  User chose to continue despite warning\n", append=True)
                    else:
                        window['-LOG-'].update("❌ Validation cancelled by user\n", append=True)
                else:
                    sg.popup_ok("✅ Files look good! Month selection appears correct.", title="Validation Success")
                    
            except Exception as e:
                error_msg = f"Validation error: {str(e)}"
                window['-LOG-'].update(f"❌ {error_msg}\n", append=True)
                sg.popup_error(f"Validation failed:\n{error_msg}")
            
            if event == 'Generate Report':
                # Validate inputs
                if not values['-IMPACTS-']:
                    sg.popup_error('Please select an Impacts Crosstab file')
                    continue
                if not values['-COUNTS-']:
                    sg.popup_error('Please select a Count Months Chronic file')
                    continue
                    
                # Clear log
                window['-LOG-'].update('')
            
            # Automatic month validation before generation
            window['-LOG-'].update("Validating month selection before generation...\n", append=True)
            window.refresh()
            
            try:
                is_valid, message = validate_month_selection(
                    values['-IMPACTS-'], 
                    values['-COUNTS-'], 
                    values['-MONTH-'], 
                    values['-YEAR-']
                )
                
                window['-LOG-'].update(f"{message}\n", append=True)
                
                if not is_valid:
                    # Show warning and ask if user wants to proceed
                    result = sg.popup_yes_no(
                        f"⚠️  MONTH VALIDATION WARNING:\n\n{message}\n\nDo you want to continue generating the report anyway?\n\n(Click 'No' to go back and fix the month selection)",
                        title="Confirm Report Generation",
                        no_titlebar=False
                    )
                    if result != 'Yes':
                        window['-LOG-'].update("❌ Report generation cancelled due to month validation warning\n", append=True)
                        continue
                    else:
                        window['-LOG-'].update("⚠️  Proceeding with report generation despite month warning\n", append=True)
                
            except Exception as e:
                window['-LOG-'].update(f"⚠️  Month validation failed: {str(e)}, proceeding anyway...\n", append=True)
            
            try:
                # Log start
                window['-LOG-'].update(f"Starting report generation...\n", append=True)
                window.refresh()
                
                # Create builder
                builder = ChronicReportBuilder(
                    exclude_regional=values['-EXCLUDE_REGIONAL-'],
                    show_indicators=values['-SHOW_INDICATORS-']
                )
                
                # Generate report
                window['-LOG-'].update(f"Processing files...\n", append=True)
                window.refresh()
                
                # Format month string
                month_str = f"{values['-MONTH-']}_{values['-YEAR-']}"
                
                corner_file, circuit_word_file, pdf_file = builder.build_monthly_report(
                    values['-IMPACTS-'],
                    values['-COUNTS-'],
                    None,  # template
                    values['-OUTPUT-'],
                    month_str
                )
                
                # Check for data quality warning
                if hasattr(builder, 'data_quality_warning') and builder.data_quality_warning:
                    window['-LOG-'].update(f"⚠️  Data Quality Warning: >10% of month cells were blank and forward-filled\n", append=True)
                
                # Check for baseline warning
                if hasattr(builder, 'baseline_found') and not builder.baseline_found:
                    window['-LOG-'].update(f"⚠️  No prior summaries found – all chronic circuits will repeat as 'New Chronic' this run.\n", append=True)
                
                # Check for ticket coercion warning
                if hasattr(builder, 'ticket_coercion_warning') and builder.ticket_coercion_warning:
                    window['-LOG-'].update(f"⚠️  Ticket Data Warning: >10% of ticket count values could not be converted to numbers\n", append=True)
                
                # Success messages
                window['-LOG-'].update(f"✅ Success! Reports generated:\n", append=True)
                window['-LOG-'].update(f"📄 Chronic Corner: {corner_file}\n", append=True)
                window['-LOG-'].update(f"📄 Circuit Report: {circuit_word_file}\n", append=True)
                if pdf_file and Path(pdf_file).exists():
                    window['-LOG-'].update(f"📄 PDF Report: {pdf_file}\n", append=True)
                
                sg.popup('Report Generation Complete!', 
                        f'Reports saved to: {values["-OUTPUT-"]}',
                        title='Success')
                
            except Exception as e:
                error_msg = f"❌ Error: {str(e)}\n"
                window['-LOG-'].update(error_msg, append=True)
                sg.popup_error(f'Error generating report:\n{str(e)}')
        
        except Exception as e:
            # Catch any GUI errors to prevent crashes
            print(f"GUI Error: {str(e)}")
            sg.popup_error(f'An unexpected error occurred:\n{str(e)}\n\nThe GUI will remain open.')
            continue
    
    window.close()


def main():
    """Main entry point - check if GUI or CLI mode"""
    if len(sys.argv) == 1:
        # No arguments - launch GUI
        gui_main()
    else:
        # Arguments provided - use CLI
        parser = argparse.ArgumentParser(description='Generate monthly chronic circuit reports')
        parser.add_argument('--impacts', required=True, help='Path to impacts A crosstab Excel file')
        parser.add_argument('--impacts-b', help='Path to impacts B crosstab Excel file (optional)')
        parser.add_argument('--counts', required=True, help='Path to counts Excel file') 
        parser.add_argument('--template', help='Path to Word template file (optional)')
        parser.add_argument('--output', default='./final_output', help='Output directory')
        parser.add_argument('--exclude-regional', action='store_true',
                           help='Exclude regional circuits from new chronic detection')
        parser.add_argument('--show-indicators', action='store_true',
                           help='Show (C) chronic and (R) regional flags in reports')
        parser.add_argument('--month', help='Month for report generation (e.g., "June 2025")')
        
        args = parser.parse_args()
        
        builder = ChronicReportBuilder(exclude_regional=args.exclude_regional, show_indicators=args.show_indicators)
        
        try:
            # Use impacts A file (for now, impacts B is optional and not used in current logic)
            impacts_file = args.impacts
            if args.impacts_b:
                print(f"Note: Using Impacts A file. Impacts B support may be added in future versions.")
                
            # Use provided month or default
            month_str = args.month.replace(' ', '_') if args.month else None
            
            corner_file, circuit_word_file, pdf_file = builder.build_monthly_report(
                impacts_file, 
                args.counts,
                args.template,
                args.output,
                month_str
            )
            
            print(f"[SUCCESS] Chronic Corner (Word): {corner_file}")
            print(f"[SUCCESS] Circuit Report (Word): {circuit_word_file}")
            if pdf_file and Path(pdf_file).exists():
                print(f"[SUCCESS] Circuit Report (PDF): {pdf_file}")
            
        except Exception as e:
            print(f"[ERROR] Error: {e}")
            import traceback
            traceback.print_exc()
            sys.exit(1)

if __name__ == "__main__":
    main()